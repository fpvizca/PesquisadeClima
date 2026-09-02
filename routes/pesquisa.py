import json
import statistics
from flask import render_template, request, redirect, session, url_for, flash, jsonify, send_file
from auth import login_required, has_role
from db import get_db
import ollama_helper

from flask import make_response

def init_routes(app):

    @app.route('/pesquisa')
    @login_required
    def pesquisa():
        db = get_db()
        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        cookie_name = f'clima_respondeu_{ciclo["id"]}'
        ja_respondeu_cookie = request.cookies.get(cookie_name)

        ja_respondeu = ja_respondeu_cookie == '1'

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
            total_perguntas = db.execute(
                "SELECT COUNT(*) as c FROM perguntas WHERE secao_id IN (SELECT id FROM secoes WHERE formulario_id = ? AND ativo = 1)",
                (ciclo['formulario_id'],)
            ).fetchone()['c']
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()
            total_perguntas = db.execute(
                "SELECT COUNT(*) as c FROM perguntas WHERE secao_id IN (SELECT id FROM secoes WHERE ativo = 1)"
            ).fetchone()['c']

        return render_template('pesquisa.html',
            ciclo=ciclo,
            ja_respondeu=ja_respondeu,
            secoes=secoes,
            total_perguntas=total_perguntas
        )

    @app.route('/pesquisa/secao/<int:secao_id>', methods=['GET', 'POST'])
    @login_required
    def pesquisa_secao(secao_id):
        db = get_db()
        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()

        secao_atual = None
        for i, s in enumerate(secoes):
            if s['id'] == secao_id:
                secao_atual = s
                secao_index = i
                break

        if not secao_atual:
            return redirect(url_for('pesquisa'))

        perguntas = db.execute(
            "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 ORDER BY ordem",
            (secao_id,)
        ).fetchall()

        # Anonymous: no respostas_existentes (can't load previous answers without usuario_id)
        respostas_existentes = {}

        if request.method == 'POST':
            for p in perguntas:
                valor = request.form.get(f'pergunta_{p["id"]}', '').strip()
                comentario = request.form.get(f'comentario_{p["id"]}', '').strip()

                # Grid: collect per-row answers into a single value
                if p['tipo'] == 'grid' and not valor:
                    grid_rows = [r.strip() for r in (p['grid_rows'] or '').split(',') if r.strip()]
                    row_vals = []
                    for i in range(len(grid_rows)):
                        rv = request.form.get(f'grid_{p["id"]}_{i}', '').strip()
                        row_vals.append(f'{grid_rows[i]}: {rv}' if rv else f'{grid_rows[i]}:')
                    valor = ' | '.join(row_vals)

                if p['obrigatoria'] and not valor:
                    flash(f'Por favor, responda: {p["texto"]}', 'danger')
                    return redirect(url_for('pesquisa_secao', secao_id=secao_id))

                if valor or comentario:
                    db.execute("""
                        INSERT INTO respostas (ciclo_id, pergunta_id, valor, comentario)
                        VALUES (?, ?, ?, ?)
                    """, (ciclo['id'], p['id'], valor if valor else None, comentario if comentario else None))
            db.commit()

            proximo_index = secao_index + 1
            if proximo_index < len(secoes):
                return redirect(url_for('pesquisa_secao', secao_id=secoes[proximo_index]['id']))
            else:
                flash('Pesquisa respondida com sucesso!', 'success')
                resp = make_response(redirect(url_for('pesquisa')))
                cookie_name = f'clima_respondeu_{ciclo["id"]}'
                resp.set_cookie(cookie_name, '1', max_age=60*60*24*365)
                return resp

        proxima_secao = secoes[secao_index + 1]['id'] if secao_index + 1 < len(secoes) else None

        return render_template('pesquisa_secao.html',
            ciclo=ciclo,
            secao=secao_atual,
            perguntas=perguntas,
            respostas_existentes=respostas_existentes,
            total_secoes=len(secoes),
            secao_atual_index=secao_index + 1,
            proxima_secao=proxima_secao,
            secoes=secoes
        )

    @app.route('/minhas-respostas')
    @login_required
    def minhas_respostas():
        flash('As respostas são anônimas e não podem ser visualizadas individualmente.', 'info')
        return redirect(url_for('pesquisa'))

    @app.route('/admin/resultados')
    @login_required
    def admin_resultados():
        db = get_db()
        usuario_id = session['usuario_id']
        if not has_role(usuario_id, 'admin'):
            flash('Acesso negado.', 'danger')
            return redirect(url_for('pesquisa'))

        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()

        dados = []
        for secao in secoes:
            perguntas = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 ORDER BY ordem",
                (secao['id'],)
            ).fetchall()
            perguntas_com_stats = []
            for p in perguntas:
                if p['tipo'] == 'escala':
                    stats = db.execute("""
                        SELECT
                            COUNT(*) as total,
                            AVG(CASE WHEN valor = 'Concordo totalmente' THEN 5
                                     WHEN valor = 'Concordo' THEN 4
                                     WHEN valor = 'Não concordo e nem discordo' THEN 3
                                     WHEN valor = 'Discordo' THEN 2
                                     WHEN valor = 'Discordo totalmente' THEN 1
                                     ELSE NULL END) as media
                        FROM respostas WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    """, (ciclo['id'], p['id'])).fetchone()
                    perguntas_com_stats.append({
                        'pergunta': p,
                        'total': stats['total'],
                        'media': round(stats['media'], 2) if stats['media'] else None
                    })
                else:
                    respostas = db.execute(
                        "SELECT valor FROM respostas WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL",
                        (ciclo['id'], p['id'])
                    ).fetchall()
                    perguntas_com_stats.append({
                        'pergunta': p,
                        'total': len(respostas),
                        'respostas': [r['valor'] for r in respostas]
                    })
            dados.append({'secao': secao, 'perguntas': perguntas_com_stats})

        total_habilitados = db.execute(
            "SELECT COUNT(*) as c FROM usuario_roles WHERE role = 'colaborador'"
        ).fetchone()['c']

        total_respostas = db.execute(
            "SELECT COUNT(*) as c FROM respostas WHERE ciclo_id = ?",
            (ciclo['id'],)
        ).fetchone()['c']

        return render_template('admin_resultados.html',
            ciclo=ciclo,
            dados=dados,
            total_habilitados=total_habilitados,
            total_respostas=total_respostas
        )

    @app.route('/admin/analise')
    @login_required
    def admin_analise():
        db = get_db()
        usuario_id = session['usuario_id']
        if not has_role(usuario_id, 'admin'):
            flash('Acesso negado.', 'danger')
            return redirect(url_for('pesquisa'))

        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()

        total_habilitados = db.execute(
            "SELECT COUNT(*) as c FROM usuario_roles WHERE role = 'colaborador'"
        ).fetchone()['c']

        total_respostas_db = db.execute(
            "SELECT COUNT(*) as c FROM respostas WHERE ciclo_id = ?",
            (ciclo['id'],)
        ).fetchone()['c']

        dados_secoes = []
        for secao in secoes:
            perguntas = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 AND tipo = 'escala' ORDER BY ordem",
                (secao['id'],)
            ).fetchall()

            medias = []
            distribuicao = {'Concordo totalmente': 0, 'Concordo': 0, 'Não concordo e nem discordo': 0, 'Discordo': 0, 'Discordo totalmente': 0}
            total_pond = 0
            total_respostas = 0

            perguntas_dados = []
            for p in perguntas:
                contagem = db.execute("""
                    SELECT valor, COUNT(*) as cnt FROM respostas
                    WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    GROUP BY valor
                """, (ciclo['id'], p['id'])).fetchall()

                mapa = {r['valor']: r['cnt'] for r in contagem}
                total_p = sum(mapa.values())

                if total_p > 0:
                    nota = (mapa.get('Concordo totalmente', 0) * 5 + mapa.get('Concordo', 0) * 4 +
                            mapa.get('Não concordo e nem discordo', 0) * 3 + mapa.get('Discordo', 0) * 2 +
                            mapa.get('Discordo totalmente', 0) * 1) / total_p
                    medias.append(nota)

                    pct_satisfatorio = ((mapa.get('Concordo totalmente', 0) + mapa.get('Concordo', 0)) / total_p * 100) if total_p > 0 else 0

                    for k in distribuicao:
                        distribuicao[k] += mapa.get(k, 0)
                    total_pond += total_p
                    total_respostas += total_p
                else:
                    nota = None
                    pct_satisfatorio = 0

                perguntas_dados.append({
                    'codigo': p['codigo'],
                    'texto': p['texto'],
                    'total': total_p,
                    'nota': round(nota, 2) if nota else None,
                    'pct_satisfatorio': round(pct_satisfatorio, 1),
                    'distribuicao': {k: mapa.get(k, 0) for k in ['Concordo totalmente', 'Concordo', 'Não concordo e nem discordo', 'Discordo', 'Discordo totalmente']}
                })

            media_secao = round(statistics.mean(medias), 2) if medias else None
            desvio = round(statistics.stdev(medias), 2) if len(medias) > 1 else 0
            pct_total_satisf = round((distribuicao['Concordo totalmente'] + distribuicao['Concordo']) / total_pond * 100, 1) if total_pond > 0 else 0

            dados_secoes.append({
                'secao': {'id': secao['id'], 'nome': secao['nome']},
                'media': media_secao,
                'desvio': desvio,
                'pct_satisfatorio': pct_total_satisf,
                'total_respostas': total_pond,
                'perguntas': perguntas_dados,
                'distribuicao': distribuicao
            })

        perguntas_abertas = db.execute("""
            SELECT p.codigo, p.texto, p.secao_id, s.nome as secao_nome, r.valor
            FROM respostas r JOIN perguntas p ON r.pergunta_id = p.id JOIN secoes s ON p.secao_id = s.id
            WHERE r.ciclo_id = ? AND p.tipo IN ('texto', 'paragrafo') AND r.valor IS NOT NULL AND r.valor != ''
            ORDER BY s.ordem, p.ordem
        """, (ciclo['id'],)).fetchall()

        return render_template('admin_analise.html',
            ciclo=ciclo,
            dados_secoes=dados_secoes,
            total_habilitados=total_habilitados,
            total_respostas=total_respostas_db,
            perguntas_abertas=perguntas_abertas
        )

    @app.route('/api/analise-ia', methods=['POST'])
    @login_required
    def api_analise_ia():
        if not has_role(session['usuario_id'], 'admin'):
            return jsonify({'error': 'Acesso negado'}), 403

        data = request.get_json()
        secao_nome = data.get('secao', 'Geral')
        dados = data.get('dados', {})

        prompt = f"""Gere uma análise profissional no estilo de relatório de pesquisa de clima organizacional para a seção "{secao_nome}".

Dados estatísticos:
- Média da seção: {dados.get('media', 'N/A')}
- Total de respostas: {dados.get('total', 'N/A')}
- % Satisfatório (Concordo + Concordo Totalmente): {dados.get('pct_satisfatorio', 'N/A')}%

Distribuição das respostas:
{json.dumps(dados.get('distribuicao', {}), ensure_ascii=False, indent=2)}

Perguntas com menores notas:
{json.dumps(dados.get('menores_notas', []), ensure_ascii=False, indent=2)}

Perguntas com maiores notas:
{json.dumps(dados.get('maiores_notas', []), ensure_ascii=False, indent=2)}

Respostas abertas relevantes:
{json.dumps(dados.get('respostas_abertas', []), ensure_ascii=False, indent=2)}

Escreva uma análise:
1. Resumo executivo da seção
2. Pontos fortes identificados
3. Pontos de atenção
4. Comparação com benchmarks (ideal > 4.0)
5. Recomendações de ações

Seja objetivo, use dados numéricos e escreva em português brasileiro profissional."""

        response = ollama_helper.generate(prompt)
        return jsonify({'response': response})

    @app.route('/admin/exportar-excel')
    @login_required
    def admin_exportar_excel():
        from flask import send_file
        import io
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        db = get_db()
        if not has_role(session['usuario_id'], 'admin'):
            flash('Acesso negado.', 'danger')
            return redirect(url_for('pesquisa'))

        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        wb = Workbook()

        # Estilos
        header_font = Font(name='Arial', bold=True, color='FFFFFF', size=11)
        header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
        section_font = Font(name='Arial', bold=True, size=12, color='1F4E79')
        normal_font = Font(name='Arial', size=10)
        green_fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
        yellow_fill = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')
        red_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        # Aba 1: Resumo por Seção
        ws_resumo = wb.active
        ws_resumo.title = 'Resumo'
        ws_resumo.column_dimensions['A'].width = 35
        ws_resumo.column_dimensions['B'].width = 15
        ws_resumo.column_dimensions['C'].width = 18
        ws_resumo.column_dimensions['D'].width = 20

        ws_resumo.merge_cells('A1:D1')
        ws_resumo['A1'] = ciclo['nome']
        ws_resumo['A1'].font = Font(name='Arial', bold=True, size=14, color='1F4E79')

        total_habilitados = db.execute(
            "SELECT COUNT(*) as c FROM usuario_roles WHERE role = 'colaborador'"
        ).fetchone()['c']

        total_respostas = db.execute(
            "SELECT COUNT(*) as c FROM respostas WHERE ciclo_id = ?",
            (ciclo['id'],)
        ).fetchone()['c']

        ws_resumo['A2'] = f'Habilitados: {total_habilitados} | Respostas: {total_respostas}'
        ws_resumo['A2'].font = Font(name='Arial', size=10, italic=True)

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()

        row = 5
        for secao in secoes:
            perguntas = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 AND tipo = 'escala' ORDER BY ordem",
                (secao['id'],)
            ).fetchall()

            medias = []
            total_pond = 0
            total_satisf = 0

            for p in perguntas:
                contagem = db.execute("""
                    SELECT valor, COUNT(*) as cnt FROM respostas
                    WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    GROUP BY valor
                """, (ciclo['id'], p['id'])).fetchall()

                mapa = {r['valor']: r['cnt'] for r in contagem}
                total_p = sum(mapa.values())
                if total_p > 0:
                    nota = (mapa.get('Concordo totalmente', 0) * 5 + mapa.get('Concordo', 0) * 4 +
                            mapa.get('Não concordo e nem discordo', 0) * 3 + mapa.get('Discordo', 0) * 2 +
                            mapa.get('Discordo totalmente', 0) * 1) / total_p
                    medias.append(nota)
                    total_pond += total_p
                    total_satisf += mapa.get('Concordo totalmente', 0) + mapa.get('Concordo', 0)

            media = round(sum(medias) / len(medias), 2) if medias else None
            pct = round(total_satisf / total_pond * 100, 1) if total_pond > 0 else 0

            ws_resumo.cell(row=row, column=1, value=secao['nome']).font = normal_font
            ws_resumo.cell(row=row, column=1).border = thin_border

            cell_media = ws_resumo.cell(row=row, column=2, value=media)
            cell_media.font = normal_font
            cell_media.alignment = Alignment(horizontal='center')
            cell_media.border = thin_border
            if media:
                cell_media.fill = green_fill if media >= 4 else yellow_fill if media >= 3 else red_fill

            cell_pct = ws_resumo.cell(row=row, column=3, value=f'{pct}%')
            cell_pct.font = normal_font
            cell_pct.alignment = Alignment(horizontal='center')
            cell_pct.border = thin_border

            ws_resumo.cell(row=row, column=4, value=total_pond).font = normal_font
            ws_resumo.cell(row=row, column=4).alignment = Alignment(horizontal='center')
            ws_resumo.cell(row=row, column=4).border = thin_border

            row += 1

        # Aba 2: Detalhamento por Pergunta
        ws_detalhe = wb.create_sheet('Detalhamento')
        ws_detalhe.column_dimensions['A'].width = 10
        ws_detalhe.column_dimensions['B'].width = 60
        ws_detalhe.column_dimensions['C'].width = 12
        ws_detalhe.column_dimensions['D'].width = 12
        ws_detalhe.column_dimensions['E'].width = 12
        ws_detalhe.column_dimensions['F'].width = 12
        ws_detalhe.column_dimensions['G'].width = 12
        ws_detalhe.column_dimensions['H'].width = 12
        ws_detalhe.column_dimensions['I'].width = 12

        row = 1
        headers2 = ['Código', 'Pergunta', 'Total', 'Média', 'Conc. Totalmente', 'Concordo', 'Neutro', 'Discordo', 'Disc. Totalmente']
        for col, h in enumerate(headers2, 1):
            cell = ws_detalhe.cell(row=row, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = thin_border

        row = 2
        for secao in secoes:
            ws_detalhe.merge_cells(start_row=row, start_column=1, end_row=row, end_column=9)
            cell = ws_detalhe.cell(row=row, column=1, value=secao['nome'])
            cell.font = section_font
            row += 1

            perguntas = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 ORDER BY ordem",
                (secao['id'],)
            ).fetchall()

            for p in perguntas:
                contagem = db.execute("""
                    SELECT valor, COUNT(*) as cnt FROM respostas
                    WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    GROUP BY valor
                """, (ciclo['id'], p['id'])).fetchall()

                mapa = {r['valor']: r['cnt'] for r in contagem}
                total_p = sum(mapa.values())

                if total_p > 0 and p['tipo'] == 'escala':
                    nota = (mapa.get('Concordo totalmente', 0) * 5 + mapa.get('Concordo', 0) * 4 +
                            mapa.get('Não concordo e nem discordo', 0) * 3 + mapa.get('Discordo', 0) * 2 +
                            mapa.get('Discordo totalmente', 0) * 1) / total_p
                    nota = round(nota, 2)
                else:
                    nota = None

                ws_detalhe.cell(row=row, column=1, value=p['codigo']).font = normal_font
                ws_detalhe.cell(row=row, column=1).border = thin_border
                ws_detalhe.cell(row=row, column=2, value=p['texto']).font = normal_font
                ws_detalhe.cell(row=row, column=2).border = thin_border
                ws_detalhe.cell(row=row, column=3, value=total_p).font = normal_font
                ws_detalhe.cell(row=row, column=3).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=3).border = thin_border

                cell_nota = ws_detalhe.cell(row=row, column=4, value=nota)
                cell_nota.font = normal_font
                cell_nota.alignment = Alignment(horizontal='center')
                cell_nota.border = thin_border
                if nota:
                    cell_nota.fill = green_fill if nota >= 4 else yellow_fill if nota >= 3 else red_fill

                ws_detalhe.cell(row=row, column=5, value=mapa.get('Concordo totalmente', 0)).font = normal_font
                ws_detalhe.cell(row=row, column=5).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=5).border = thin_border

                ws_detalhe.cell(row=row, column=6, value=mapa.get('Concordo', 0)).font = normal_font
                ws_detalhe.cell(row=row, column=6).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=6).border = thin_border

                ws_detalhe.cell(row=row, column=7, value=mapa.get('Não concordo e nem discordo', 0)).font = normal_font
                ws_detalhe.cell(row=row, column=7).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=7).border = thin_border

                ws_detalhe.cell(row=row, column=8, value=mapa.get('Discordo', 0)).font = normal_font
                ws_detalhe.cell(row=row, column=8).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=8).border = thin_border

                ws_detalhe.cell(row=row, column=9, value=mapa.get('Discordo totalmente', 0)).font = normal_font
                ws_detalhe.cell(row=row, column=9).alignment = Alignment(horizontal='center')
                ws_detalhe.cell(row=row, column=9).border = thin_border

                row += 1

        # Aba 3: Respostas Abertas
        perguntas_abertas = db.execute("""
            SELECT p.codigo, p.texto, s.nome as secao_nome, r.valor
            FROM respostas r JOIN perguntas p ON r.pergunta_id = p.id JOIN secoes s ON p.secao_id = s.id
            WHERE r.ciclo_id = ? AND p.tipo IN ('texto', 'paragrafo') AND r.valor IS NOT NULL AND r.valor != ''
            ORDER BY s.ordem, p.ordem
        """, (ciclo['id'],)).fetchall()

        if perguntas_abertas:
            ws_abertas = wb.create_sheet('Respostas Abertas')
            ws_abertas.column_dimensions['A'].width = 10
            ws_abertas.column_dimensions['B'].width = 50
            ws_abertas.column_dimensions['C'].width = 30
            ws_abertas.column_dimensions['D'].width = 60

            row = 1
            headers3 = ['Código', 'Pergunta', 'Seção', 'Resposta']
            for col, h in enumerate(headers3, 1):
                cell = ws_abertas.cell(row=row, column=col, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', wrap_text=True)
                cell.border = thin_border

            row = 2
            for r in perguntas_abertas:
                ws_abertas.cell(row=row, column=1, value=r['codigo']).font = normal_font
                ws_abertas.cell(row=row, column=1).border = thin_border
                ws_abertas.cell(row=row, column=2, value=r['texto']).font = normal_font
                ws_abertas.cell(row=row, column=2).border = thin_border
                ws_abertas.cell(row=row, column=3, value=r['secao_nome']).font = normal_font
                ws_abertas.cell(row=row, column=3).border = thin_border
                ws_abertas.cell(row=row, column=4, value=r['valor']).font = normal_font
                ws_abertas.cell(row=row, column=4).border = thin_border
                ws_abertas.cell(row=row, column=4).alignment = Alignment(wrap_text=True)
                row += 1

        # Salvar
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f'Resultados_Clima_{ciclo["ano"]}.xlsx'
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True, download_name=filename)

    @app.route('/admin/exportar-word')
    @login_required
    def admin_exportar_word():
        import io
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        db = get_db()
        if not has_role(session['usuario_id'], 'admin'):
            flash('Acesso negado.', 'danger')
            return redirect(url_for('pesquisa'))

        ciclo = db.execute("SELECT * FROM ciclos WHERE ativo = 1 ORDER BY id DESC LIMIT 1").fetchone()
        if not ciclo:
            flash('Nenhum ciclo ativo encontrado.', 'warning')
            return redirect(url_for('index'))

        total_habilitados = db.execute(
            "SELECT COUNT(*) as c FROM usuario_roles WHERE role = 'colaborador'"
        ).fetchone()['c']
        total_respostas = db.execute(
            "SELECT COUNT(*) as c FROM respostas WHERE ciclo_id = ?",
            (ciclo['id'],)
        ).fetchone()['c']

        # Fetch sections linked to the cycle's form
        if ciclo['formulario_id']:
            secoes = db.execute(
                "SELECT * FROM secoes WHERE formulario_id = ? AND ativo = 1 ORDER BY ordem",
                (ciclo['formulario_id'],)
            ).fetchall()
        else:
            secoes = db.execute("SELECT * FROM secoes WHERE ativo = 1 ORDER BY ordem").fetchall()

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # -- Capa --
        for _ in range(6):
            doc.add_paragraph('')

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Pesquisa de Clima Organizacional')
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ciclo['nome'])
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph('')

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Relatório de Resultados')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Ano: {ciclo["ano"]}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

        # -- Sumário Executivo --
        h = doc.add_heading('1. Sumário Executivo', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        doc.add_paragraph(
            f'O presente relatório apresenta os resultados da {ciclo["nome"]}, '
            f'realizada no ano de {ciclo["ano"]}. A pesquisa teve como objetivo analisar a percepção '
            f'dos colaboradores em relação ao ambiente de trabalho, cultura organizacional, '
            f'bem-estar, diversidade e inclusão.'
        )

        p = doc.add_paragraph()
        run = p.add_run('Participação: ')
        run.bold = True
        p.add_run(f'{total_respostas} respostas de {total_habilitados} colaboradores habilitados '
                  f'({round(total_respostas / total_habilitados * 100, 1) if total_habilitados > 0 else 0}% de adesão)')

        # -- Resumo por Seção --
        h = doc.add_heading('2. Resumo por Seção', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, txt in enumerate(['Seção', 'Média', '% Satisfatório', 'Respostas']):
            hdr[i].text = txt
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        dados_resumo = []
        for secao in secoes:
            perguntas_esc = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 AND tipo = 'escala' ORDER BY ordem",
                (secao['id'],)
            ).fetchall()

            medias = []
            total_pond = 0
            total_satisf = 0

            for p in perguntas_esc:
                contagem = db.execute("""
                    SELECT valor, COUNT(*) as cnt FROM respostas
                    WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    GROUP BY valor
                """, (ciclo['id'], p['id'])).fetchall()

                mapa = {r['valor']: r['cnt'] for r in contagem}
                total_p = sum(mapa.values())
                if total_p > 0:
                    nota = (mapa.get('Concordo totalmente', 0) * 5 + mapa.get('Concordo', 0) * 4 +
                            mapa.get('Não concordo e nem discordo', 0) * 3 + mapa.get('Discordo', 0) * 2 +
                            mapa.get('Discordo totalmente', 0) * 1) / total_p
                    medias.append(nota)
                    total_pond += total_p
                    total_satisf += mapa.get('Concordo totalmente', 0) + mapa.get('Concordo', 0)

            media = round(sum(medias) / len(medias), 2) if medias else None
            pct = round(total_satisf / total_pond * 100, 1) if total_pond > 0 else 0

            row = table.add_row().cells
            row[0].text = secao['nome']
            row[1].text = str(media) if media else '-'
            row[2].text = f'{pct}%'
            row[3].text = str(total_pond)
            dados_resumo.append({'nome': secao['nome'], 'media': media, 'pct': pct})

        doc.add_paragraph('')

        # -- Gráficos Resumo --
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.ticker as mticker
        import tempfile
        import os

        tmp_dir = tempfile.mkdtemp()

        nomes_secoes = [d['nome'] for d in dados_resumo if d['media'] is not None]
        medias_secoes = [d['media'] for d in dados_resumo if d['media'] is not None]
        pcts_secoes = [d['pct'] for d in dados_resumo if d['media'] is not None]

        if medias_secoes:
            h = doc.add_heading('2.1 Visão Geral', level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(31, 78, 121)

            # Gráfico de Médias por Seção
            fig, ax = plt.subplots(figsize=(8, 4))
            colors = ['#2ecc71' if m >= 4 else '#f39c12' if m >= 3 else '#e74c3c' for m in medias_secoes]
            bars = ax.barh(nomes_secoes, medias_secoes, color=colors, edgecolor='white', height=0.6)
            ax.set_xlim(0, 5)
            ax.set_xlabel('Média (escala 1 a 5)', fontsize=10)
            ax.set_title('Média por Seção', fontsize=13, fontweight='bold', color='#1F4E79', pad=12)
            ax.xaxis.set_major_locator(mticker.FixedLocator([1, 2, 3, 4, 5]))
            for bar, val in zip(bars, medias_secoes):
                ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2, f'{val}', va='center', fontsize=9, fontweight='bold')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            plt.tight_layout()
            path_media = os.path.join(tmp_dir, 'media_secoes.png')
            fig.savefig(path_media, dpi=200, bbox_inches='tight')
            plt.close(fig)

            doc.add_picture(path_media, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run('Gráfico 1 – Média de satisfação por seção (escala 1 a 5)')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

            doc.add_paragraph('')

            # Gráfico de % Satisfatório
            fig2, ax2 = plt.subplots(figsize=(8, 4))
            colors2 = ['#2ecc71' if p >= 70 else '#f39c12' if p >= 50 else '#e74c3c' for p in pcts_secoes]
            bars2 = ax2.barh(nomes_secoes, pcts_secoes, color=colors2, edgecolor='white', height=0.6)
            ax2.set_xlim(0, 100)
            ax2.set_xlabel('% Satisfatório', fontsize=10)
            ax2.set_title('Percentual Satisfatório por Seção', fontsize=13, fontweight='bold', color='#1F4E79', pad=12)
            for bar, val in zip(bars2, pcts_secoes):
                ax2.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, f'{val}%', va='center', fontsize=9, fontweight='bold')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            plt.tight_layout()
            path_pct = os.path.join(tmp_dir, 'pct_secoes.png')
            fig2.savefig(path_pct, dpi=200, bbox_inches='tight')
            plt.close(fig2)

            doc.add_picture(path_pct, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap2 = doc.add_paragraph()
            p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap2.add_run('Gráfico 2 – Percentual de respostas satisfatórias (Concordo + Concordo totalmente)')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

            doc.add_page_break()

        # -- Detalhamento por Seção --
        h = doc.add_heading('3. Detalhamento por Seção', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        for secao in secoes:
            h2 = doc.add_heading(secao['nome'], level=2)
            for run in h2.runs:
                run.font.color.rgb = RGBColor(31, 78, 121)

            if secao['descricao']:
                p = doc.add_paragraph()
                run = p.add_run(secao['descricao'].replace('<p class="mb-0">', '').replace('<p class="mb-2">', '').replace('<hr class="my-2">', '---').replace('<strong>', '').replace('</strong>', '').replace('<em>', '').replace('</em>', '').replace('<br>', '\n').replace('</p>', '').strip())
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.italic = True

            perguntas = db.execute(
                "SELECT * FROM perguntas WHERE secao_id = ? AND ativo = 1 ORDER BY ordem",
                (secao['id'],)
            ).fetchall()

            for p in perguntas:
                contagem = db.execute("""
                    SELECT valor, COUNT(*) as cnt FROM respostas
                    WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL
                    GROUP BY valor
                """, (ciclo['id'], p['id'])).fetchall()

                mapa = {r['valor']: r['cnt'] for r in contagem}
                total_p = sum(mapa.values())

                if p['tipo'] == 'escala' and total_p > 0:
                    nota = (mapa.get('Concordo totalmente', 0) * 5 + mapa.get('Concordo', 0) * 4 +
                            mapa.get('Não concordo e nem discordo', 0) * 3 + mapa.get('Discordo', 0) * 2 +
                            mapa.get('Discordo totalmente', 0) * 1) / total_p
                    pct_satisf = round((mapa.get('Concordo totalmente', 0) + mapa.get('Concordo', 0)) / total_p * 100, 1)

                    p_doc = doc.add_paragraph()
                    run = p_doc.add_run(f'{p["codigo"]} - {p["texto"]}')
                    run.bold = True
                    run.font.size = Pt(10)

                    tbl = doc.add_table(rows=2, cols=6)
                    tbl.style = 'Light List Accent 1'
                    headers_esc = ['', 'Concordo totalmente', 'Concordo', 'Neutro', 'Discordo', 'Discordo totalmente']
                    for i, txt in enumerate(headers_esc):
                        tbl.rows[0].cells[i].text = txt
                        for par in tbl.rows[0].cells[i].paragraphs:
                            for r in par.runs:
                                r.font.size = Pt(9)
                                r.bold = True

                    vals = [total_p, mapa.get('Concordo totalmente', 0), mapa.get('Concordo', 0),
                            mapa.get('Não concordo e nem discordo', 0), mapa.get('Discordo', 0),
                            mapa.get('Discordo totalmente', 0)]
                    for i, v in enumerate(vals):
                        tbl.rows[1].cells[i].text = str(v)
                        for par in tbl.rows[1].cells[i].paragraphs:
                            for r in par.runs:
                                r.font.size = Pt(9)

                    p_nota = doc.add_paragraph()
                    run = p_nota.add_run(f'Média: {round(nota, 2)} | Satisfatório: {pct_satisf}%')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(31, 78, 121)

                    # Gráfico de pizza para a pergunta
                    labels = ['Concordo totalmente', 'Concordo', 'Neutro', 'Discordo', 'Discordo totalmente']
                    sizes = [mapa.get(l, 0) for l in labels]
                    if sum(sizes) > 0:
                        colors_pie = ['#27ae60', '#2ecc71', '#f39c12', '#e74c3c', '#c0392b']
                        fig_pie, ax_pie = plt.subplots(figsize=(5, 3))
                        wedges, texts, autotexts = ax_pie.pie(
                            [s for s in sizes if s > 0],
                            labels=[l for l, s in zip(labels, sizes) if s > 0],
                            colors=[c for c, s in zip(colors_pie, sizes) if s > 0],
                            autopct='%1.0f%%', startangle=90, textprops={'fontsize': 8}
                        )
                        for at in autotexts:
                            at.set_fontsize(8)
                            at.set_fontweight('bold')
                        ax_pie.set_title(f'{p["codigo"]}', fontsize=9, fontweight='bold', color='#1F4E79', pad=8)
                        plt.tight_layout()
                        path_pie = os.path.join(tmp_dir, f'q_{p["id"]}.png')
                        fig_pie.savefig(path_pie, dpi=150, bbox_inches='tight')
                        plt.close(fig_pie)

                        doc.add_picture(path_pie, width=Inches(3.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap_q = doc.add_paragraph()
                        p_cap_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p_cap_q.add_run(f'Fig. – Distribuição de respostas: {p["texto"][:80]}')
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                        run.italic = True

                elif p['tipo'] in ('texto', 'paragrafo'):
                    respostas_texto = db.execute("""
                        SELECT valor FROM respostas
                        WHERE ciclo_id = ? AND pergunta_id = ? AND valor IS NOT NULL AND valor != ''
                    """, (ciclo['id'], p['id'])).fetchall()

                    if respostas_texto:
                        p_doc = doc.add_paragraph()
                        run = p_doc.add_run(f'{p["codigo"]} - {p["texto"]}')
                        run.bold = True
                        run.font.size = Pt(10)

                        for resp in respostas_texto:
                            p_resp = doc.add_paragraph(style='List Bullet')
                            run = p_resp.add_run(resp['valor'])
                            run.font.size = Pt(10)

                doc.add_paragraph('')

        # -- Respostas Abertas --
        perguntas_abertas = db.execute("""
            SELECT p.codigo, p.texto, s.nome as secao_nome, r.valor
            FROM respostas r JOIN perguntas p ON r.pergunta_id = p.id JOIN secoes s ON p.secao_id = s.id
            WHERE r.ciclo_id = ? AND p.tipo IN ('texto', 'paragrafo') AND r.valor IS NOT NULL AND r.valor != ''
            ORDER BY s.ordem, p.ordem
        """, (ciclo['id'],)).fetchall()

        if perguntas_abertas:
            h = doc.add_heading('4. Respostas Abertas', level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(31, 78, 121)

            current_secao = None
            for resp in perguntas_abertas:
                if resp['secao_nome'] != current_secao:
                    current_secao = resp['secao_nome']
                    h2 = doc.add_heading(current_secao, level=2)
                    for run in h2.runs:
                        run.font.color.rgb = RGBColor(31, 78, 121)

                p_doc = doc.add_paragraph()
                run = p_doc.add_run(f'{resp["codigo"]} - {resp["texto"]}')
                run.bold = True
                run.font.size = Pt(10)

                p_resp = doc.add_paragraph(style='List Bullet')
                run = p_resp.add_run(resp['valor'])
                run.font.size = Pt(10)

            doc.add_paragraph('')

        # -- Considerações Finais --
        h = doc.add_heading('Considerações Finais', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        melhor_secao = max(dados_resumo, key=lambda x: x['media'] if x['media'] else 0) if dados_resumo else None
        pior_secao = min(dados_resumo, key=lambda x: x['media'] if x['media'] else 999) if dados_resumo else None

        if melhor_secao and pior_secao:
            doc.add_paragraph(
                f'A seção com melhor avaliação foi "{melhor_secao["nome"]}" '
                f'(média {melhor_secao["media"]}), enquanto "{pior_secao["nome"]}" '
                f'apresentou a menor média ({pior_secao["media"]}).'
            )

        doc.add_paragraph(
            'Os dados coletados servem como base para o desenvolvimento de planos de ação '
            'que visam melhorar continuamente o ambiente de trabalho e a experiência dos colaboradores.'
        )

        # Limpar arquivos temporários
        import shutil
        try:
            shutil.rmtree(tmp_dir)
        except Exception:
            pass

        # Salvar
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f'Relatorio_Clima_{ciclo["ano"]}.docx'
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
